"""Processador de documentos multi-formato (PDF, DOCx, TXT, HTML)."""

from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
import re

from app.domain.value_objects.document_metadata import DocumentMetadata, ChunkMetadata

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processa documentos de diferentes formatos e os converte em chunks para indexação.

    Formatos suportados:
    - PDF (.pdf)
    - Word (.docx)
    - Excel (.xlsx) - opcional
    - Texto (.txt)
    - HTML (.html)
    """

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Args:
            chunk_size: Tamanho máximo de cada chunk em caracteres
            chunk_overlap: Sobreposição entre chunks para manter contexto
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """
        Extrai texto de um arquivo PDF.

        Args:
            pdf_path: Caminho para o arquivo PDF

        Returns:
            Texto extraído
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(pdf_path)
            full_text = ""

            for page_num, page in enumerate(doc):
                text = page.get_text()
                full_text += f"\n{text}\n"
                logger.debug(f"Página {page_num + 1}/{len(doc)} extraída")

            doc.close()
            logger.info(f"PDF extraído: {pdf_path.name} ({len(full_text)} caracteres)")
            return full_text

        except Exception as e:
            logger.error(f"Erro ao extrair PDF {pdf_path}: {e}")
            raise

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """
        Extrai texto de um arquivo Word (.docx).

        Args:
            docx_path: Caminho para o arquivo DOCX

        Returns:
            Texto extraído
        """
        try:
            from docx import Document

            doc = Document(docx_path)
            full_text = []

            # Extrair parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Extrair tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text.append(row_text)

            text = "\n".join(full_text)
            logger.info(f"DOCX extraído: {docx_path.name} ({len(text)} caracteres)")
            return text

        except Exception as e:
            logger.error(f"Erro ao extrair DOCX {docx_path}: {e}")
            raise

    def extract_text_from_txt(self, txt_path: Path) -> str:
        """
        Extrai texto de arquivo TXT.

        Args:
            txt_path: Caminho para o arquivo TXT

        Returns:
            Texto extraído
        """
        try:
            # Tentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"TXT extraído: {txt_path.name} (encoding: {encoding})")
                    return text
                except UnicodeDecodeError:
                    continue

            raise ValueError(f"Não foi possível decodificar {txt_path} com encodings testados")

        except Exception as e:
            logger.error(f"Erro ao ler TXT {txt_path}: {e}")
            raise

    def extract_text_from_html(self, html_path: Path) -> str:
        """
        Extrai texto de arquivo HTML.

        Args:
            html_path: Caminho para o arquivo HTML

        Returns:
            Texto extraído (sem tags HTML)
        """
        try:
            from bs4 import BeautifulSoup

            with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remover scripts e styles
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text(separator='\n')

            # Limpar linhas vazias
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            text = '\n'.join(lines)

            logger.info(f"HTML extraído: {html_path.name} ({len(text)} caracteres)")
            return text

        except Exception as e:
            logger.error(f"Erro ao extrair HTML {html_path}: {e}")
            raise

    def extract_text(self, file_path: Path) -> str:
        """
        Extrai texto de um arquivo baseado em sua extensão.

        Args:
            file_path: Caminho para o arquivo

        Returns:
            Texto extraído

        Raises:
            ValueError: Se formato não for suportado
        """
        suffix = file_path.suffix.lower()

        extractors = {
            '.pdf': self.extract_text_from_pdf,
            '.docx': self.extract_text_from_docx,
            '.txt': self.extract_text_from_txt,
            '.html': self.extract_text_from_html,
            '.htm': self.extract_text_from_html,
        }

        extractor = extractors.get(suffix)
        if not extractor:
            raise ValueError(
                f"Formato não suportado: {suffix}. "
                f"Formatos suportados: {list(extractors.keys())}"
            )

        return extractor(file_path)

    def chunk_text(self, text: str) -> List[str]:
        """
        Divide texto em chunks com overlap para manter contexto.

        Args:
            text: Texto completo do documento

        Returns:
            Lista de chunks de texto
        """
        # Limpar texto
        text = self._clean_text(text)

        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # Se não é o último chunk, tentar quebrar em parágrafo/sentença
            if end < len(text):
                # Procurar quebra de parágrafo
                paragraph_break = text.rfind('\n\n', start, end)
                if paragraph_break != -1 and paragraph_break > start:
                    end = paragraph_break

                # Se não achou parágrafo, procurar quebra de sentença
                elif end < len(text):
                    sentence_breaks = ['. ', '! ', '? ', '.\n']
                    for sep in sentence_breaks:
                        sentence_break = text.rfind(sep, start, end)
                        if sentence_break != -1 and sentence_break > start:
                            end = sentence_break + len(sep)
                            break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Próximo chunk com overlap
            start = end - self.chunk_overlap if end < len(text) else end

        logger.debug(f"Texto dividido em {len(chunks)} chunks")
        return chunks

    def _clean_text(self, text: str) -> str:
        """
        Limpa texto removendo espaços excessivos e caracteres especiais.

        Args:
            text: Texto bruto

        Returns:
            Texto limpo
        """
        # Remover múltiplos espaços
        text = re.sub(r' +', ' ', text)

        # Remover múltiplas quebras de linha (manter no máximo 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remover espaços no início/fim de cada linha
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    def process_document(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> List[ChunkMetadata]:
        """
        Processa um documento completo e retorna chunks com metadados.

        Args:
            file_path: Caminho para o arquivo
            metadata: Metadados do documento

        Returns:
            Lista de ChunkMetadata prontos para indexação
        """
        logger.info(f"Processando documento: {file_path.name}")

        # Extrair texto
        full_text = self.extract_text(file_path)

        if not full_text or len(full_text) < 10:
            logger.warning(f"Documento vazio ou muito curto: {file_path.name}")
            return []

        # Dividir em chunks
        chunks = self.chunk_text(full_text)

        if not chunks:
            logger.warning(f"Nenhum chunk gerado para: {file_path.name}")
            return []

        # Criar ChunkMetadata para cada chunk
        chunk_metadata_list = []
        for idx, chunk_text in enumerate(chunks):
            chunk_meta = ChunkMetadata.from_document_metadata(
                doc_metadata=metadata,
                chunk_index=idx,
                total_chunks=len(chunks),
                text=chunk_text
            )
            chunk_metadata_list.append(chunk_meta)

        logger.info(
            f"Documento processado: {file_path.name} "
            f"→ {len(chunks)} chunks ({len(full_text)} caracteres)"
        )

        return chunk_metadata_list

    def get_document_stats(self, file_path: Path) -> Dict[str, Any]:
        """
        Retorna estatísticas sobre um documento sem processá-lo completamente.

        Args:
            file_path: Caminho para o arquivo

        Returns:
            Dict com estatísticas (tamanho, número de páginas, etc.)
        """
        stats = {
            "file_name": file_path.name,
            "file_size_bytes": file_path.stat().st_size,
            "file_format": file_path.suffix,
        }

        try:
            if file_path.suffix.lower() == '.pdf':
                import fitz
                doc = fitz.open(file_path)
                stats["pages"] = len(doc)
                doc.close()

            elif file_path.suffix.lower() == '.docx':
                from docx import Document
                doc = Document(file_path)
                stats["paragraphs"] = len(doc.paragraphs)
                stats["tables"] = len(doc.tables)

        except Exception as e:
            logger.warning(f"Erro ao obter stats de {file_path.name}: {e}")

        return stats
